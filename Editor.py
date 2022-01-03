import os
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client as win32
from shutil import copyfile

def readTxtFile(path, dic):
    file = open(path, "r", encoding='UTF-8')
    
    while True:
        line = file.readline()
        if not line:
            break
        
        raw_res = line.strip()
        pair = raw_res.split(':')
        
        if pair[0] == '':
            continue
        
        # print(pair[0], pair[1])
        dic[pair[0]] = pair[1]

    file.close()
    
def readExcelFile(path, dic):
    file = open(path, "r", encoding='UTF-8')
    
    while True:
        line = file.readline()
        if not line:
            break
        
        raw_res = line.strip()
        pair = raw_res.split(':')
        
        # print(pair[0], pair[1])
        dic[pair[0]] = pair[1]

    file.close()

def editHwpFiles(name, dic):
    filenames = os.listdir(os.getcwd())
    fileformat = '.hwp'

    for filename in filenames:
        full_filename = os.path.join(name, filename)
        ext = os.path.splitext(full_filename)[-1]
        
        if (ext == '.hwp'):
            # print(full_filename)
            
            hwp = win32.gencache.EnsureDispatch("HWPFrame.HwpObject")
            hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
            
            editname = full_filename.replace(fileformat, '')
            resultname = editname + '_수정본' + fileformat
            
            copyfile(full_filename, resultname)
            
            hwp.Open(resultname)
            hwp.InitScan()
            
            for i in dic:
                hwp.HAction.GetDefault("AllReplace", hwp.HParameterSet.HFindReplace.HSet)
                option=hwp.HParameterSet.HFindReplace
                option.FindString = i
                option.ReplaceString = dic[i]
                option.ReplaceCharShape.TextColor = hwp.RGBColor(255, 0, 0)
                #option.ReplaceCharShape.Bold = 1
                option.IgnoreMessage = 1
                hwp.HAction.Execute("AllReplace", hwp.HParameterSet.HFindReplace.HSet)
                    
            hwp.ReleaseScan()
            hwp.Clear(3)
            hwp.Quit()
    
def editWordFiles(name, dic):
    filenames = os.listdir(os.getcwd())
    fileformat = '.docx'

    for filename in filenames:
        full_filename = os.path.join(name, filename)
        ext = os.path.splitext(full_filename)[-1]
        
        if (ext == fileformat):
            # print(full_filename)
            document = Document(full_filename)

            paragraphs = list(document.paragraphs)
            for p in paragraphs:
                for key, val in dic.items():
                    if key in p.text:
                        run = p.runs
                        started = False
                        key_index = 0
                        found_runs = list()
                        found_all = False
                        
                        for i in range(len(run)):
                            
                            # case 2: search for partial text, find first run
                            if key[key_index] in run[i].text and run[i].text[-1] in key and not started:
                                # check sequence
                                start_index = run[i].text.find(key[key_index])
                                check_length = len(run[i].text)
                                
                                #print("1:" + run[i].text)
                                #print("2:" + key[key_index])
                                
                                for text_index in range(start_index, check_length):
                                    if run[i].text[text_index] != key[key_index]:
                                        # no match so must be false positive
                                        break
                                if key_index == 0:
                                    started = True
                                
                                chars_found = check_length - start_index
                                
                                if len(key) < chars_found:
                                    chars_found = len(key)
                                
                                key_index += chars_found
                                found_runs.append((i, start_index, chars_found))
                                
                                if key_index != len(key):
                                    continue
                                else:
                                    # found all chars in key_name
                                    found_all = True
                                    break

                            # case 2: search for partial text, find subsequent run
                            if key[key_index] in run[i].text and started and not found_all:
                                # check sequence
                                chars_found = 0
                                check_length = len(run[i].text)
                                
                                for text_index in range(0, check_length):
                                    
                                    if len(key) < key_index + 1:
                                        key_index = len(key) - 1
                                    
                                    if run[i].text[text_index] == key[key_index]:
                                        key_index += 1
                                        chars_found += 1
                                    else:
                                        break
                                # no match so must be end
                                found_runs.append((i, 0, chars_found))
                                if key_index == len(key):
                                    found_all = True
                                    break

                        if found_all:
                            for i, item in enumerate(found_runs):
                                index, start, length = [t for t in item]
                                if i == 0:
                                    text = run[index].text.replace(run[index].text[start:start + length], str(val))
                                    run[index].text = text
                                    run[index].font.color.rgb = RGBColor(255, 0, 0)
                                else:
                                    text = run[index].text.replace(run[index].text[start:start + length], '')
                                    run[index].text = text
                                    run[index].font.color.rgb = RGBColor(255, 0, 0)
                        # print(p.text)
                        
            editname = filename.replace(fileformat, '')
            resultname = editname + '_수정본' + fileformat
            
            document.save(resultname)
            
def editRtfFiles(name, dic):
    filenames = os.listdir(os.getcwd())
    fileformat = '.rtf'

    for filename in filenames:
        full_filename = os.path.join(name, filename)
        ext = os.path.splitext(full_filename)[-1]
        
        if (ext == fileformat):
            word = win32.Dispatch("Word.Application")
            doc = word.Documents.Open(full_filename)
            
            editname = full_filename.replace(fileformat, '')
            resultname = editname + '_수정본' + '.docx'
            
            wdFormatDocumentDefault = 16
            
            doc.SaveAs(resultname, FileFormat=wdFormatDocumentDefault)
            doc.Close()
            word.Quit()
            
            document = Document(resultname)
            
            paragraphs = list(document.paragraphs)
            for p in paragraphs:
                for key, val in dic.items():
                    if key in p.text:
                        run = p.runs
                        started = False
                        key_index = 0
                        found_runs = list()
                        found_all = False
                        
                        for i in range(len(run)):
                            
                            # case 2: search for partial text, find first run
                            if key[key_index] in run[i].text and run[i].text[-1] in key and not started:
                                # check sequence
                                start_index = run[i].text.find(key[key_index])
                                check_length = len(run[i].text)
                                
                                #print("1:" + run[i].text)
                                #print("2:" + key[key_index])
                                
                                for text_index in range(start_index, check_length):
                                    if run[i].text[text_index] != key[key_index]:
                                        # no match so must be false positive
                                        break
                                if key_index == 0:
                                    started = True
                                
                                chars_found = check_length - start_index
                                
                                if len(key) < chars_found:
                                    chars_found = len(key)
                                
                                key_index += chars_found
                                found_runs.append((i, start_index, chars_found))
                                
                                if key_index != len(key):
                                    continue
                                else:
                                    # found all chars in key_name
                                    found_all = True
                                    break

                            # case 2: search for partial text, find subsequent run
                            if key[key_index] in run[i].text and started and not found_all:
                                # check sequence
                                chars_found = 0
                                check_length = len(run[i].text)
                                
                                for text_index in range(0, check_length):
                                    
                                    if len(key) < key_index + 1:
                                        key_index = len(key) - 1
                                    
                                    if run[i].text[text_index] == key[key_index]:
                                        key_index += 1
                                        chars_found += 1
                                    else:
                                        break
                                # no match so must be end
                                found_runs.append((i, 0, chars_found))
                                if key_index == len(key):
                                    found_all = True
                                    break

                        if found_all:
                            for i, item in enumerate(found_runs):
                                index, start, length = [t for t in item]
                                if i == 0:
                                    text = run[index].text.replace(run[index].text[start:start + length], str(val))
                                    run[index].text = text
                                    run[index].font.color.rgb = RGBColor(255, 0, 0)
                                else:
                                    text = run[index].text.replace(run[index].text[start:start + length], '')
                                    run[index].text = text
                                    run[index].font.color.rgb = RGBColor(255, 0, 0)
                        # print(p.text)
            
            document.save(resultname)
            os.remove(resultname)
            
            editname = full_filename.replace(fileformat, '')
            resultname = editname + '_수정본' + '.rtf'
            document.save(resultname)

if __name__ == "__main__":
    
    myDictionary = {}
    dirname = os.getcwd()
    notepath = dirname + '\\note.txt'
    
    readTxtFile(notepath, myDictionary)
    #readExcelFile(notepath, myDictionary)
    editWordFiles(dirname, myDictionary)
    editHwpFiles(dirname, myDictionary)
    editRtfFiles(dirname, myDictionary)